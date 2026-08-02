from docx import Document
from datetime import datetime
import re

from services.timetable_importer import TimetableImporter


class DocxTimetableImporter:

    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.helper = TimetableImporter(filepath)

    def preview_tables(self):

        output = []

        for table_no, table in enumerate(self.doc.tables, start=1):

            output.append(f"\n========== TABLE {table_no} ==========\n")

            for row in table.rows:

                cells = [cell.text.strip() for cell in row.cells]

                output.append(" | ".join(cells))

        return "\n".join(output)
    def import_docx(self):

        records = []

        for table in self.doc.tables:

            # Skip header row
            for row in table.rows[1:]:

                cells = [cell.text.strip() for cell in row.cells]

                # Ignore incomplete rows
                if len(cells) < 8:
                    continue

                print(cells)

                # ---------------------------------
                # DAY + DATE
                # ---------------------------------
                text = cells[0].replace("\n", " ")
                text = " ".join(text.split())

                day = ""

                if "SATURDAY" in text.upper():
                    day = "Saturday"
                elif "SUNDAY" in text.upper():
                    day = "Sunday"

                exam_date = None

                match = re.search(r"\d{2}/\d{2}/\d{4}", text)

                if match:
                    exam_date = datetime.strptime(
                        match.group(),
                        "%d/%m/%Y"
                    ).date()

                # ---------------------------------
                # TIME
                # ---------------------------------
                time_text = cells[1]

                lines = [
                    x.strip()
                    for x in time_text.splitlines()
                    if x.strip()
                ]

                if len(lines) >= 2:

                    start_time = datetime.strptime(
                        lines[0],
                        "%I:%M%p"
                    ).time()

                    end_time = datetime.strptime(
                        lines[1],
                        "%I:%M%p"
                    ).time()

                else:
                    start_time = None
                    end_time = None

                # ---------------------------------
                # CLASS
                # ---------------------------------
                class_name = cells[2]

                programme, level = self.helper.extract_class_details(
                    class_name
                )

                # ---------------------------------
                # COURSE
                # ---------------------------------
                course_code = cells[3]
                course_title = cells[4]

                students = cells[5]

                venue = cells[6]

                examiner = cells[7]

                # ---------------------------------
                # SAVE
                # ---------------------------------
                records.append({

                    "day": day,
                    "exam_date": exam_date,

                    "start_time": start_time,
                    "end_time": end_time,

                    "programme": programme,
                    "level": level,

                    "session": "Weekend",

                    "class": class_name,

                    "course_code": course_code,
                    "course_title": course_title,

                    "students": students,

                    "venue": venue,
                    "examiner": examiner

                })

        print("=" * 80)
        print("TOTAL DOCX RECORDS:", len(records))

        for r in records[:10]:
            print(r)

        print("=" * 80)

        return records